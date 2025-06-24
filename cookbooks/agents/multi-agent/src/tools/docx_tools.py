import os
from datetime import datetime
from src.tools.common import client
from docx import Document

from src.tools.common import judgment
@judgment.observe_tools()
class DocxTools:
    """Flow-layout document editing tools using Word documents (.docx) with automatic reflow."""

    def create_docx(self, filename: str = "report") -> str:
        """Input: base filename without extension (str) | Action: create new empty Word document with timestamp | Output: full .docx file path (str)"""
    
        os.makedirs("reports", exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = f"reports/{filename}_{timestamp}.docx"
        
        doc = Document()
        doc.add_paragraph()
        doc.save(docx_path)
        
        print(f"[AGENT] 📄 Word document created: {docx_path}")
        return f"Word document created: {docx_path}"
      

    def read_docx(self, docx_path: str) -> str:
        """Input: path to existing .docx file (str) | Action: read document like a human viewing it | Output: structured content with elements (str)"""
        try:
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            doc = Document(docx_path)
            content = f"WORD DOCUMENT: {docx_path}\n{'='*60}\n\n"
            
            element_count = 0
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    element_count += 1

                    style = paragraph.style.name
                    content += f"PARAGRAPH {i+1} | Style: {style}\n"
                    content += f"Text: '{text}'\n"
                    
                    for run in paragraph.runs:
                        if run.bold or run.italic or run.underline:
                            formatting = []
                            if run.bold: formatting.append("Bold")
                            if run.italic: formatting.append("Italic") 
                            if run.underline: formatting.append("Underline")
                            content += f"  Formatting: {', '.join(formatting)}\n"
                    content += "\n"
            
            for i, table in enumerate(doc.tables):
                element_count += 1
                content += f"TABLE {i+1} | Rows: {len(table.rows)} | Columns: {len(table.columns)}\n"
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    content += f"  Row {row_idx+1}: {row_text}\n"
                content += "\n"
            
            inline_shapes = len(doc.inline_shapes)
            if inline_shapes > 0:
                content += f"IMAGES/SHAPES: {inline_shapes} embedded objects\n\n"
                element_count += inline_shapes
            
            content += f"TOTAL ELEMENTS: {element_count} (paragraphs, tables, images)\n"
            return content
                
        except Exception as e:
            return f"Error reading Word document: {str(e)}"

    def add_text_to_docx(self, docx_path: str, content: str, style: str = "Normal", position: str = "end") -> str:
        """Input: .docx path + text + style + position | Action: add text that automatically flows with existing content | Output: success confirmation (str)"""
        try:
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            doc = Document(docx_path)
            
            
            if position == "end":
                
                paragraph = doc.add_paragraph(content)
            elif position == "beginning":
               
                new_para = doc.paragraphs[0]._element
                new_para.getparent().insert(0, new_para)
                paragraph = doc.paragraphs[0]
                paragraph.text = content
            else:
                
                paragraph = doc.add_paragraph(content)
            
         
            try:
                paragraph.style = style
            except:
                paragraph.style = "Normal"
            
            doc.save(docx_path)
            
            print(f"[AGENT] ✏️  Text added to document: '{content[:50]}...'")
            return f"Text added at {position}: '{content[:50]}...'"
            
        except Exception as e:
            return f"Error adding text: {str(e)}"

    def add_heading_to_docx(self, docx_path: str, heading_text: str, level: int = 1) -> str:
        """Input: .docx path + heading text + level (1-9) | Action: add heading that automatically formats and flows | Output: success confirmation (str)"""
        try:
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            doc = Document(docx_path)
            
            # Add heading (level 0 = Title, 1-9 = Heading levels)
            heading = doc.add_heading(heading_text, level=level)
            doc.save(docx_path)
            
            print(f"[AGENT] 📋 Heading added: {heading_text}")
            return f"Level {level} heading added: '{heading_text}'"
            
        except Exception as e:
            return f"Error adding heading: {str(e)}"

    def add_image_to_docx(self, docx_path: str, image_path: str, width_inches: float = None, caption: str = None) -> str:
        """Input: .docx path + image path + optional width + caption | Action: insert image that flows with text automatically | Output: success confirmation (str)"""
        try:
            from docx.shared import Inches
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            if not os.path.exists(image_path):
                return f"Error: Image file not found: {image_path}"
            
            doc = Document(docx_path)
            
            if width_inches:
                doc.add_picture(image_path, width=Inches(width_inches))
            else:
                doc.add_picture(image_path)
            
            if caption:
                caption_para = doc.add_paragraph(caption)
                caption_para.style = "Caption"
            
            doc.save(docx_path)
            
            print(f"[AGENT] 🖼️  Image added: {os.path.basename(image_path)}")
            return f"Image added: {image_path}" + (f" with caption: '{caption}'" if caption else "")
            
        except Exception as e:
            return f"Error adding image: {str(e)}"

    def add_table_to_docx(self, docx_path: str, data: list, headers: bool = True) -> str:
        """Input: .docx path + table data (list of lists) + has headers | Action: insert table that flows with document | Output: success confirmation (str)"""
        try:
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            if not data or not isinstance(data, list):
                return f"Error: Invalid table data provided"
            
            doc = Document(docx_path)
            
            rows = len(data)
            cols = len(data[0]) if data else 1
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            for row_idx, row_data in enumerate(data):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        table.rows[row_idx].cells[col_idx].text = str(cell_data)
            
            if headers and len(table.rows) > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            
            doc.save(docx_path)
            
            print(f"[AGENT] 📊 Table added: {rows}x{cols}")
            return f"Table added: {rows} rows x {cols} columns"
            
        except Exception as e:
            return f"Error adding table: {str(e)}"

    def delete_text_from_docx(self, docx_path: str, text_to_delete: str) -> str:
        """Input: .docx path + specific text to delete | Action: find and remove specific text from document | Output: deletion confirmation (str)"""
        try:
            from docx import Document
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            doc = Document(docx_path)
            deletions_made = 0
            
            for paragraph in doc.paragraphs:
                if text_to_delete in paragraph.text:
                    original_text = paragraph.text
                    paragraph.text = paragraph.text.replace(text_to_delete, "")
                    if original_text != paragraph.text:
                        deletions_made += 1
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if text_to_delete in paragraph.text:
                                original_text = paragraph.text
                                paragraph.text = paragraph.text.replace(text_to_delete, "")
                                if original_text != paragraph.text:
                                    deletions_made += 1
            
            doc.save(docx_path)
            
            if deletions_made > 0:
                print(f"[AGENT] 🗑️  Deleted {deletions_made} instances of '{text_to_delete}'")
                return f"Deleted {deletions_made} instances of '{text_to_delete}' from document"
            else:
                return f"Text '{text_to_delete}' not found in document"
            
        except Exception as e:
            return f"Error deleting text: {str(e)}"

    def delete_paragraph_from_docx(self, docx_path: str, paragraph_number: int) -> str:
        """Input: .docx path + paragraph number (1-based) | Action: delete specific paragraph | Output: deletion confirmation (str)"""
        try:
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            doc = Document(docx_path)
            
            if paragraph_number < 1 or paragraph_number > len(doc.paragraphs):
                return f"Error: Paragraph {paragraph_number} doesn't exist (document has {len(doc.paragraphs)} paragraphs)"
            
            # Delete paragraph (convert to 0-based index)
            paragraph = doc.paragraphs[paragraph_number - 1]
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
            
            doc.save(docx_path)
            
            print(f"[AGENT] 🗑️  Paragraph {paragraph_number} deleted")
            return f"Paragraph {paragraph_number} deleted from document"
            
        except Exception as e:
            return f"Error deleting paragraph: {str(e)}"

    def add_page_break_to_docx(self, docx_path: str) -> str:
        """Input: .docx path | Action: insert page break to start new page | Output: success confirmation (str)"""
        try:
            from docx import Document
            from docx.enum.text import WD_BREAK
            
            if not os.path.exists(docx_path):
                return f"Error: Word document not found: {docx_path}"
            
            doc = Document(docx_path)
            
            paragraph = doc.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            
            doc.save(docx_path)
            
            print(f"[AGENT] 📄 Page break added")
            return f"Page break added to document"
            
        except Exception as e:
            return f"Error adding page break: {str(e)}"
   
     